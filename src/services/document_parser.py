"""
Document parser for the Semantic Plagiarism Detector
Extracts text from various document formats.
"""

import os
import io
import re
from pathlib import Path
from typing import Optional, Tuple, Dict, Any
import logging
from datetime import datetime

from src.config.settings import settings
from src.models.document import Document, DocumentType, DocumentStatus
from src.utils.file_validators import validate_file

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Parses various document formats and extracts text.
    Supports: TXT, PDF, DOCX, DOC, RTF, ODT
    """
    
    def __init__(self):
        self.supported_formats = {
            '.txt': self._parse_txt,
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_doc,
            '.rtf': self._parse_rtf,
            '.odt': self._parse_odt,
        }
        self.max_extract_length = settings.MAX_EXTRACT_LENGTH
        self.min_text_length = settings.MIN_TEXT_LENGTH
        self._available_parsers = self._check_parser_availability()
    
    def _check_parser_availability(self) -> Dict[str, bool]:
        """Check which parsers are available."""
        available = {}
        
        try:
            import PyPDF2  # noqa
            available['pdf'] = True
        except ImportError:
            available['pdf'] = False
        
        try:
            import docx  # noqa
            available['docx'] = True
        except ImportError:
            available['docx'] = False
        
        try:
            import striprtf  # noqa
            available['rtf'] = True
        except ImportError:
            available['rtf'] = False
        
        try:
            import odf  # noqa
            available['odt'] = True
        except ImportError:
            available['odt'] = False
        
        return available
    
    def parse_document(self, file_path: Path) -> Document:
        """Parse a document and extract text."""
        validation = validate_file(file_path)
        if not validation.is_valid:
            doc = Document(
                filename=file_path.name,
                original_filename=file_path.name,
                file_path=str(file_path),
                file_size=file_path.stat().st_size if file_path.exists() else 0,
                status=DocumentStatus.FAILED,
                error_message=validation.error_message
            )
            return doc
        
        extension = file_path.suffix.lower()
        file_type = self._get_file_type(extension)
        
        doc = Document(
            filename=file_path.name,
            original_filename=file_path.name,
            file_path=str(file_path),
            file_size=file_path.stat().st_size,
            file_type=file_type,
            mime_type=validation.mime_type,
            status=DocumentStatus.PROCESSING
        )
        
        doc.generate_hash()
        
        try:
            parser_method = self.supported_formats.get(extension)
            if parser_method:
                content, metadata = parser_method(file_path)
                doc.content = content[:self.max_extract_length] if content else ""
                doc.metadata = metadata or {}
                doc._update_content_stats()
                doc.status = DocumentStatus.COMPLETED
                doc.processed_at = datetime.now()
            else:
                doc.status = DocumentStatus.FAILED
                doc.error_message = f"Unsupported file format: {extension}"
        except Exception as e:
            logger.error(f"Failed to parse document {file_path}: {e}")
            doc.status = DocumentStatus.FAILED
            doc.error_message = str(e)
        
        return doc
    
    def _get_file_type(self, extension: str) -> DocumentType:
        mapping = {
            '.txt': DocumentType.TXT,
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.doc': DocumentType.DOC,
            '.rtf': DocumentType.RTF,
            '.odt': DocumentType.ODT,
        }
        return mapping.get(extension, DocumentType.UNKNOWN)
    
    def _parse_txt(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                with open(file_path, 'rb') as f:
                    content = f.read().decode('utf-8', errors='replace')
            
            content = self._clean_text(content)
            
            metadata = {
                'encoding': encoding if content else 'unknown',
                'line_count': content.count('\n') + 1 if content else 0
            }
            
            return content or "", metadata
        except Exception as e:
            logger.error(f"Failed to parse TXT file {file_path}: {e}")
            raise
    
    def _parse_pdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        try:
            if not self._available_parsers.get('pdf', False):
                raise ImportError("PyPDF2 is not installed. Install with: pip install PyPDF2")
            
            import PyPDF2
            
            content_parts = []
            metadata = {}
            
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                if reader.metadata:
                    for key, value in reader.metadata.items():
                        metadata[key] = str(value) if value else ''
                
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        content_parts.append(f"[Page {page_num + 1}]\n{text}")
            
            content = "\n\n".join(content_parts)
            content = self._clean_text(content)
            
            metadata['page_count'] = len(reader.pages)
            metadata['pdf_version'] = str(reader.pdf_header) if hasattr(reader, 'pdf_header') else ''
            
            return content or "", metadata
        except ImportError as e:
            logger.error(f"PDF parsing failed: {e}")
            raise
        except Exception as e:
            logger.error(f"Failed to parse PDF file {file_path}: {e}")
            raise
    
    def _parse_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        try:
            if not self._available_parsers.get('docx', False):
                raise ImportError("python-docx is not installed. Install with: pip install python-docx")
            
            import docx
            
            doc = docx.Document(file_path)
            
            content_parts = []
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
            
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    if any(row_text):
                        content_parts.append(" | ".join(row_text))
            
            content = "\n".join(content_parts)
            content = self._clean_text(content)
            
            try:
                if doc.core_properties:
                    for prop in ['author', 'title', 'subject', 'created', 'modified']:
                        if hasattr(doc.core_properties, prop):
                            val = getattr(doc.core_properties, prop)
                            if val:
                                metadata[prop] = str(val)
            except:
                pass
            
            return content or "", metadata
        except ImportError as e:
            logger.error(f"DOCX parsing failed: {e}")
            raise
        except Exception as e:
            logger.error(f"Failed to parse DOCX file {file_path}: {e}")
            raise
    
    def _parse_doc(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        try:
            try:
                import textract
                content = textract.process(str(file_path)).decode('utf-8', errors='ignore')
                content = self._clean_text(content)
                return content or "", {'method': 'textract'}
            except ImportError:
                pass
            
            with open(file_path, 'rb') as f:
                content = f.read()
                text_pattern = re.compile(b'[\\x20-\\x7E]{4,}')
                text_bytes = text_pattern.findall(content)
                text_parts = [t.decode('ascii', errors='ignore') for t in text_bytes]
                content = "\n".join(text_parts)
                content = self._clean_text(content)
                
                if len(content) > 0:
                    return content or "", {'method': 'binary_extraction'}
                
                raise Exception("Could not extract text from DOC file")
        except Exception as e:
            logger.error(f"Failed to parse DOC file {file_path}: {e}")
            raise
    
    def _parse_rtf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        try:
            if not self._available_parsers.get('rtf', False):
                raise ImportError("striprtf is not installed. Install with: pip install striprtf")
            
            import striprtf
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()
            
            content = striprtf.rtf_to_text(rtf_content)
            content = self._clean_text(content)
            
            return content or "", {'method': 'striprtf'}
        except ImportError as e:
            logger.error(f"RTF parsing failed: {e}")
            raise
        except Exception as e:
            logger.error(f"Failed to parse RTF file {file_path}: {e}")
            raise
    
    def _parse_odt(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        try:
            if not self._available_parsers.get('odt', False):
                raise ImportError("odfpy is not installed. Install with: pip install odfpy")
            
            import odf.opendocument
            from odf import text, teletype
            
            odt_doc = odf.opendocument.load(str(file_path))
            
            content_parts = []
            
            for para in odt_doc.getElementsByType(text.P):
                content_parts.append(teletype.extractText(para))
            
            content = "\n".join(content_parts)
            content = self._clean_text(content)
            
            return content or "", {'method': 'odfpy', 'paragraph_count': len(content_parts)}
        except ImportError as e:
            logger.error(f"ODT parsing failed: {e}")
            raise
        except Exception as e:
            logger.error(f"Failed to parse ODT file {file_path}: {e}")
            raise
    
    def _clean_text(self, text: str) -> str:
        if not text:
            return ""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s.,!?;:()\-\'"]', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = text.strip()
        return text
    
    def get_document_info(self, file_path: Path) -> Dict[str, Any]:
        if not file_path.exists():
            return {'error': 'File does not exist'}
        
        stats = file_path.stat()
        
        return {
            'filename': file_path.name,
            'extension': file_path.suffix,
            'size_bytes': stats.st_size,
            'size_mb': stats.st_size / (1024 * 1024),
            'created_at': datetime.fromtimestamp(stats.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(stats.st_mtime).isoformat(),
            'is_readable': os.access(str(file_path), os.R_OK)
        }


def parse_document(file_path: Path) -> Document:
    parser = DocumentParser()
    return parser.parse_document(file_path)


def get_document_info(file_path: Path) -> Dict[str, Any]:
    parser = DocumentParser()
    return parser.get_document_info(file_path)