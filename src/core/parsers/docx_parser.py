# MIT License
#
# Copyright (c) 2026 Ganesh Kambli
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
# SOFTWARE.

"""src/core/parsers/docx_parser.py - Word document (.docx, .doc) text extraction strategy."""

import io
import logging
import os
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass

from src.core.parsers.common import PDFInput

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class ParsedDocxText:
    """Parsed DOCX text together with heading metadata for each word."""

    text: str
    headings: list[str | None]

    @property
    def word_headings(self) -> list[str | None]:
        """Backward-compatible alias for the heading metadata."""
        return self.headings


def extract_text_from_docx(file: PDFInput) -> str:
    """Extract text from a DOCX file, prefixing headings with Markdown # markers."""
    try:
        import docx

        doc_file = io.BytesIO(file) if isinstance(file, bytes) else file
        document = docx.Document(doc_file)

        current_heading = None
        word_headings = []
        paragraphs_text = []

        for paragraph in document.paragraphs:
            p_text = paragraph.text
            style_name = paragraph.style.name if paragraph.style else ""

            heading_match = re.match(r"^Heading\s+(\d+)$", style_name or "")
            if heading_match:
                level = int(heading_match.group(1))
                prefix = "#" * level + " "
                p_text = prefix + p_text
                current_heading = p_text.strip()

            paragraphs_text.append(p_text)
            p_words = p_text.split()
            word_headings.extend([current_heading] * len(p_words))

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        p_text = paragraph.text
                        paragraphs_text.append(p_text)
                        p_words = p_text.split()
                        word_headings.extend([current_heading] * len(p_words))

        full_text = "\n\n".join(paragraphs_text)
        return ParsedDocxText(full_text.strip(), word_headings=word_headings)
    except Exception as exc:
        logger.error(f"[document_parser] Error reading DOCX: {exc}")
    return ""


def extract_text_from_doc(file: PDFInput) -> str:
    """Extract plain text from a legacy Word Document (.doc) using antiword."""
    if not shutil.which("antiword"):
        logger.warning(
            "antiword binary not found. Please install antiword to parse .doc files."
        )
        raise RuntimeError(
            "antiword binary is not installed on the system. Cannot parse .doc files."
        )

    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as temp_file:
        if isinstance(file, bytes):
            temp_file.write(file)
        elif isinstance(file, str):
            with open(file, "rb") as f:
                temp_file.write(f.read())
        else:
            content = file.read()
            if isinstance(content, str):
                content = content.encode("utf-8")
            temp_file.write(content)
        temp_file_path = temp_file.name

    try:
        result = subprocess.run(
            ["antiword", temp_file_path],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
            timeout=30,
            check=True,
        )
        return result.stdout.strip()
    except subprocess.CalledProcessError as exc:
        logger.error(f"[document_parser] antiword failed: {exc.stderr}")
        raise RuntimeError(
            f"antiword failed to extract text from .doc file: {exc.stderr}"
        ) from exc
    except subprocess.TimeoutExpired as exc:
        logger.error(f"[document_parser] antiword timed out: {exc}")
        raise RuntimeError("antiword execution timed out.") from exc
    finally:
        try:
            os.remove(temp_file_path)
        except OSError:
            pass
