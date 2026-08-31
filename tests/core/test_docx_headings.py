import io

import docx

from src.core.document_parser import extract_text_from_docx
from src.core.text_chunking import chunk_text


def test_docx_headings_extraction_and_chunking():
    # Create an in-memory DOCX file
    doc = docx.Document()
    doc.add_paragraph("Heading 1 Title", style="Heading 1")
    doc.add_paragraph(
        "This is the content of the first section. It has some text to chunk."
    )
    doc.add_paragraph("Heading 2 Title", style="Heading 2")
    doc.add_paragraph(
        "This is the content of the second section. More text for chunking."
    )
    doc.add_paragraph("Heading 3 Title", style="Heading 3")
    doc.add_paragraph(
        "This is the content of the deeply nested third section. Testing hierarchy."
    )

    # Save to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_bytes = file_stream.getvalue()

    # Extract text
    parsed_text = extract_text_from_docx(file_bytes)
    assert parsed_text.text
    assert parsed_text.headings
    assert len(parsed_text.headings) == len(parsed_text.text.split())

    # Chunk text
    chunks = chunk_text(parsed_text, chunk_size=30, chunk_overlap=5)

    assert len(chunks) >= 3

    # First chunk inherits the correct section title (Heading 1 Title)
    assert chunks[0].metadata.get("section_title") == "# Heading 1 Title"

    # Heading changes updating subsequent chunks
    found_heading_2 = False
    found_heading_3 = False
    for chunk in chunks:
        if "second section" in chunk:
            assert chunk.metadata.get("section_title") == "## Heading 2 Title"
            found_heading_2 = True
        elif "deeply nested third section" in chunk:
            assert chunk.metadata.get("section_title") == "### Heading 3 Title"
            found_heading_3 = True

    assert found_heading_2
    assert found_heading_3


def test_docx_no_headings():
    doc = docx.Document()
    doc.add_paragraph("Just normal text here.")
    doc.add_paragraph("No headings at all.")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_bytes = file_stream.getvalue()

    parsed_text = extract_text_from_docx(file_bytes)
    chunks = chunk_text(parsed_text, chunk_size=30, chunk_overlap=5)

    for chunk in chunks:
        assert not chunk.metadata or chunk.metadata.get("section_title") is None


def test_non_docx_no_headings():
    # Test simple plain text string
    normal_text = "This is a simple plain text string that has no docx structure."
    chunks = chunk_text(normal_text, chunk_size=30, chunk_overlap=5)
    for chunk in chunks:
        assert (
            not hasattr(chunk, "metadata")
            or chunk.metadata.get("section_title") is None
        )


def test_extract_text_from_docx_with_tables():
    """Verify that extract_text_from_docx extracts text from tables inside DOCX documents."""
    from src.core.parsers.docx_parser import (
        extract_text_from_docx as parser_extract_docx,
    )

    # Create an in-memory DOCX file with a 2x2 table
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)

    # Fill cells with distinct text
    table.cell(0, 0).text = "Cell 1,1 Text"
    table.cell(0, 1).text = "Cell 1,2 Text"
    table.cell(1, 0).text = "Cell 2,1 Text"
    table.cell(1, 1).text = "Cell 2,2 Text"

    # Save to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_bytes = file_stream.getvalue()

    # Extract using the main document parser version
    parsed_text_1 = extract_text_from_docx(file_bytes)
    assert "Cell 1,1 Text" in parsed_text_1.text
    assert "Cell 1,2 Text" in parsed_text_1.text
    assert "Cell 2,1 Text" in parsed_text_1.text
    assert "Cell 2,2 Text" in parsed_text_1.text

    # Extract using the parsers package version
    parsed_text_2 = parser_extract_docx(file_bytes)
    assert "Cell 1,1 Text" in parsed_text_2
    assert "Cell 1,2 Text" in parsed_text_2
    assert "Cell 2,1 Text" in parsed_text_2
    assert "Cell 2,2 Text" in parsed_text_2
