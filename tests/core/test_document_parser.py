import io
import shutil

import zipfile
from unittest.mock import MagicMock, patch

import docx
import fitz  # PyMuPDF
import pytest

from src.core.document_parser import (
    CorruptedArchiveError,
    extract_text,
    extract_text_from_docx,
    extract_text_from_pdf,
    extract_text_from_txt,
    extract_text_from_zip,
    extract_texts,
    strip_bibliography,
)

import time


from src.core.document_parser import (clean_text, extract_text_from_odt,
                                     remove_ignore_phrases)

# Skip OCR tests when Tesseract binary is not present on this machine
TESSERACT_AVAILABLE = shutil.which("tesseract") is not None


def _make_pdf_bytes(text: str) -> bytes:
    """Create a minimal in-memory PDF containing the given text."""
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    # Ensure there are enough words to bypass OCR fallback (at least 8 words)
    words = (text + " word" * 10).split()
    c.drawString(50, 150, " ".join(words))
    c.showPage()
    c.save()
    return buf.getvalue()


def _make_encrypted_pdf_bytes(text: str = "Confidential Content", password: str = "secret123") -> bytes:
    """Create an in-memory password-protected (encrypted) PDF using PyMuPDF (fitz)."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), text)
    pdf_bytes = doc.tobytes(
        encryption=fitz.PDF_ENCRYPT_AES_256,
        user_pw=password,
        owner_pw="owner_pass",
    )
    doc.close()
    return pdf_bytes


def _make_docx_bytes(text: str) -> bytes:
    """Create a minimal in-memory DOCX containing the given text."""
    doc = docx.Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_odt_bytes(text: str) -> bytes:
    """Create a minimal in-memory ODT containing the given text."""
    content_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<office:document-content '
        'xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" '
        'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" '
        'office:version="1.2">'
        '<office:body>'
        '<office:text>'
        f'<text:p>{text}</text:p>'
        '</office:text>'
        '</office:body>'
        '</office:document-content>'
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("content.xml", content_xml.encode("utf-8"))
    return buf.getvalue()


def _make_valid_zip_bytes(files: dict) -> bytes:
    """Create a valid in-memory ZIP archive containing given file names and contents."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        for filename, content in files.items():
            zf.writestr(filename, content)
    return buf.getvalue()


def _make_large_docx_bytes(num_pages: int = 100) -> bytes:
    """Create a multi-page in-memory DOCX containing realistic paragraphs."""
    doc = docx.Document()
    sample_paragraph = (
        "Semantic Plagiarism Detection System performance benchmark paragraph. "
        "This paragraph simulates student submission content across multiple pages "
        "to ensure high-throughput processing and memory efficiency during analysis."
    )
    for i in range(num_pages):
        doc.add_heading(f"Chapter {i + 1}: Section Overview", level=2)
        doc.add_paragraph(f"Page {i + 1} content. {sample_paragraph}")
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_headings() -> bytes:
    """Create an in-memory DOCX with Heading 1, Heading 2, and Normal paragraphs."""
    doc = docx.Document()
    doc.add_heading("Chapter 1", level=1)
    doc.add_paragraph("Introductory paragraph.")
    doc.add_heading("Section A", level=2)
    doc.add_paragraph("Detailed content here.")
    doc.add_heading("Subsection", level=3)
    doc.add_paragraph("Even more detail.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@patch("src.core.document_parser._ocr_pdf_page", return_value="")
def test_extract_from_pdf_bytes(mock_ocr):
    pdf_bytes = _make_pdf_bytes(
        "Hello PDF this is a document with enough words to satisfy native text check"
    )

    # For blank page PDF, pdfplumber might return empty string, but it shouldn't error
    result = extract_text_from_pdf(pdf_bytes)
    assert isinstance(result, str)


def test_extract_from_pdf_filters_repeated_headers_page_numbers_and_whitespace():
    # Build mock pages where header/footer lines are REPEATED across pages
    # and page numbers sit on their own dedicated lines (so the filter strips them)
    page_one = MagicMock()
    page_one.extract_text.return_value = (
        "Research Report\n"
        "Introduction\n"
        "This section introduces the topic in detail with enough words.\n"
        "Page 1"
    )
    page_two = MagicMock()
    page_two.extract_text.return_value = (
        "Research Report\n"
        "Body content is written here at length for analysis purposes.\n"
        "Page 2"
    )

    fake_pdf = MagicMock()
    fake_pdf.pages = [page_one, page_two]
    fake_pdf.__enter__ = MagicMock(return_value=fake_pdf)
    fake_pdf.__exit__ = MagicMock(return_value=False)

    with patch("src.core.document_parser.pdfplumber.open", return_value=fake_pdf):
        result = extract_text_from_pdf(io.BytesIO(b"%PDF-fake-pdf"))

    # Repeated header across all pages must be stripped
    assert "Research Report" not in result
    # Standalone page-number lines must be stripped
    assert "Page 1" not in result
    assert "Page 2" not in result
    # Body content must survive
    assert "Introduction" in result
    assert "Body content" in result
    assert "\n\n\n" not in result


class TestEncryptedPDFHandling:
    """Test suite verifying encrypted/password-protected PDF detection in document_parser.py (#828)."""

    def test_extract_text_from_encrypted_pdf_handles_gracefully(self):
        """Encrypted PDF should handle gracefully without unhandled crashes."""
        encrypted_pdf_bytes = _make_encrypted_pdf_bytes(
            text="Protected Student Assignment Content", password="secret_password"
        )
        # Should cleanly return empty string or handled error signal without crashing
        result = extract_text_from_pdf(encrypted_pdf_bytes)
        assert isinstance(result, str)
        assert "Protected Student Assignment Content" not in result

    def test_extract_text_routing_encrypted_pdf(self):
        """Routing encrypted PDF through extract_text should return empty/handled text safely."""
        encrypted_pdf_bytes = _make_encrypted_pdf_bytes(
            text="Protected Content", password="pass"
        )
        result = extract_text(encrypted_pdf_bytes, "encrypted_submission.pdf")
        assert isinstance(result, str)


def test_extract_from_docx_bytes():
    docx_bytes = _make_docx_bytes("Hello DOCX")
    result = extract_text_from_docx(docx_bytes)
    assert result == "Hello DOCX"


def test_extract_from_odt_bytes():
    odt_bytes = _make_odt_bytes("Hello ODT")
    result = extract_text_from_odt(odt_bytes)
    assert result == "Hello ODT"


def test_docx_large_document_extraction_benchmark():
    """Benchmark test asserting 100-page DOCX extraction completes under 2.0 seconds (#579)."""
    large_docx_bytes = _make_large_docx_bytes(num_pages=100)

    start_time = time.perf_counter()
    extracted_text = extract_text_from_docx(large_docx_bytes)
    elapsed_time = time.perf_counter() - start_time

    assert len(extracted_text) > 0
    assert "Chapter 100: Section Overview" in extracted_text
    assert elapsed_time < 2.0, f"DOCX extraction took {elapsed_time:.3f}s (expected < 2.0s)"


def test_extract_text_routing_odt():
    odt_bytes = _make_odt_bytes("ODT content via routing")
    result = extract_text(odt_bytes, "test.odt")
    assert result == "ODT content via routing"


def test_extract_from_docx_heading_markers():
    docx_bytes = _make_docx_with_headings()
    result = extract_text_from_docx(docx_bytes)
    assert "# Chapter 1" in result
    assert "## Section A" in result
    assert "### Subsection" in result
    assert "Introductory paragraph." in result
    assert "Detailed content here." in result
    assert "Even more detail." in result


def test_extract_from_docx_plain_paragraph_unchanged():
    """Normal paragraphs without heading style must not get # prefixes."""
    docx_bytes = _make_docx_bytes("Just a normal paragraph.")
    result = extract_text_from_docx(docx_bytes)
    assert result == "Just a normal paragraph."
    assert not result.startswith("#")


def test_extract_text_routing_docx_with_headings():
    """Verify heading markers survive the full extract_text routing pipeline."""
    docx_bytes = _make_docx_with_headings()
    result = extract_text(docx_bytes, "test.docx")
    assert "# Chapter 1" in result
    assert "## Section A" in result


def test_extract_from_txt_bytes():
    txt_bytes = b"Hello TXT"
    result = extract_text_from_txt(txt_bytes)
    assert result == "Hello TXT"


# ---------------------------------------------------------------------------
# Corrupted Zip Submission Tests (#580)
# ---------------------------------------------------------------------------


class TestCorruptedZipHandling:

    def test_extract_text_from_valid_zip(self):
        zip_bytes = _make_valid_zip_bytes({"essay1.txt": "First student essay text.", "essay2.txt": "Second student submission."})
        result = extract_text_from_zip(zip_bytes)
        assert "First student essay text." in result
        assert "Second student submission." in result

    def test_corrupted_zip_header_raises_user_friendly_error(self):
        corrupted_bytes = b"PK\x03\x04corrupted_zip_header_data_not_valid_archive"
        with pytest.raises(CorruptedArchiveError) as exc_info:
            extract_text_from_zip(corrupted_bytes)
        assert "corrupted" in str(exc_info.value).lower()

    def test_routing_corrupted_zip_via_extract_text(self):
        corrupted_bytes = b"PK\x03\x04_corrupted_zip_header_data"
        with pytest.raises(CorruptedArchiveError):
            extract_text(corrupted_bytes, "submission_batch.zip")


@patch("src.core.document_parser._ocr_pdf_page", return_value="")
def test_extract_text_routing(mock_ocr):
    pdf_bytes = _make_pdf_bytes(
        "Hello PDF this is a document with enough words to satisfy native text check"
    )

    docx_bytes = _make_docx_bytes("Hello DOCX")
    txt_bytes = b"Hello TXT"

    assert isinstance(extract_text(pdf_bytes, "test.pdf"), str)
    assert extract_text(docx_bytes, "test.docx") == "Hello DOCX"
    assert extract_text(txt_bytes, "test.txt") == "Hello TXT"
    # Fallback case (now rejected by security check)
    assert extract_text(txt_bytes, "test.unknown") == ""


def test_extract_texts_mixed():
    docx_bytes = _make_docx_bytes("Hello DOCX")
    txt_bytes = b"Hello TXT"

    mock_file1 = MagicMock()
    mock_file1.name = "doc1.docx"
    mock_file1.read.return_value = docx_bytes

    mock_file2 = MagicMock()
    mock_file2.name = "doc2.txt"
    mock_file2.read.return_value = txt_bytes

    # Mock extract_text to isolate testing of extract_texts structure
    with patch(
        "src.core.document_parser.extract_text",
        side_effect=lambda f, name, **kwargs: f"Parsed {name}",
    ):
        results = extract_texts([mock_file1, mock_file2])

    assert results["doc1.docx"] == "Parsed doc1.docx"
    assert results["doc2.txt"] == "Parsed doc2.txt"


# ---------------------------------------------------------------------------
# strip_bibliography tests (Issue #116)
# ---------------------------------------------------------------------------


class TestStripBibliography:

    def test_strips_references_header(self):
        text = "Some body text.\n\nReferences\n[1] Smith, 2020.\n[2] Jones, 2021."
        result = strip_bibliography(text)
        assert result == "Some body text."
        assert "Smith" not in result

    def test_strips_works_cited(self):
        text = "Analysis complete.\n\nWorks Cited\nDoe, J. (2019). Paper."
        result = strip_bibliography(text)
        assert result == "Analysis complete."

    def test_strips_bibliography_header(self):
        text = "Conclusion drawn.\n\nBibliography\nAdams, B. Book."
        result = strip_bibliography(text)
        assert result == "Conclusion drawn."

    def test_strips_citations_header(self):
        text = "Findings discussed.\n\nCitations\nLee, 2018."
        result = strip_bibliography(text)
        assert result == "Findings discussed."

    def test_strips_reference_list_header(self):
        text = "Summary provided.\n\nReference List\nWang, 2022."
        result = strip_bibliography(text)
        assert result == "Summary provided."

    def test_strips_sources_header(self):
        text = "Method described.\n\nSources\nData from WHO."
        result = strip_bibliography(text)
        assert result == "Method described."

    def test_case_insensitive(self):
        text = "Body here.\n\nREFERENCES\n[1] entry."
        result = strip_bibliography(text)
        assert result == "Body here."

    def test_preserves_normal_text(self):
        text = (
            "Introduction section with enough words to be meaningful.\n\n"
            "Methodology describes the approach used in this study.\n\n"
            "Results show significant improvement over baseline.\n\n"
            "Conclusion summarizes key findings."
        )
        result = strip_bibliography(text)
        assert result == text

    def test_no_bibliography_unchanged(self):
        text = "Just a plain document with no special headers at all."
        assert strip_bibliography(text) == text

    def test_empty_string(self):
        assert strip_bibliography("") == ""

    def test_inline_references_not_stripped(self):
        text = "The references to prior work are important.\nMore text follows."
        result = strip_bibliography(text)
        assert result == text

    def test_bibliography_not_at_start_of_line_not_stripped(self):
        text = "The Bibliography section was reviewed.\nMore text."
        result = strip_bibliography(text)
        assert result == text

    def test_extract_text_strips_bibliography_from_txt(self):
        txt_bytes = b"Body text.\n\nReferences\n[1] Entry one."
        result = extract_text(txt_bytes, "test.txt")
        assert "References" not in result
        assert "Body text" in result

    def test_extract_text_strips_bibliography_from_docx(self):
        docx_bytes = _make_docx_bytes("Body content.\n\nBibliography\nEntry one.")
        result = extract_text(docx_bytes, "test.docx")
        assert "Bibliography" not in result
        assert "Body content" in result


# ---------------------------------------------------------------------------
# remove_ignore_phrases tests (Issue #161)
# ---------------------------------------------------------------------------


class TestRemoveIgnorePhrases:

    def test_removes_single_phrase(self):
        text = (
            "Q1: Explain the theory of relativity. This is my answer about relativity."
        )
        ignore_phrases = "Q1: Explain the theory of relativity"
        result = remove_ignore_phrases(text, ignore_phrases)
        assert "Q1: Explain the theory of relativity" not in result
        assert "This is my answer about relativity" in result

    def test_removes_multiple_phrases(self):
        text = "Q1: First question. My answer to first. Q2: Second question. My answer to second."
        ignore_phrases = "Q1: First question\nQ2: Second question"
        result = remove_ignore_phrases(text, ignore_phrases)
        assert "Q1: First question" not in result
        assert "Q2: Second question" not in result
        assert "My answer to first" in result
        assert "My answer to second" in result

    def test_empty_ignore_phrases_returns_original(self):
        text = "This is my original text."
        ignore_phrases = ""
        result = remove_ignore_phrases(text, ignore_phrases)
        assert result == text

    def test_whitespace_only_ignore_phrases_returns_original(self):
        text = "This is my original text."
        ignore_phrases = "   \n\n   "
        result = remove_ignore_phrases(text, ignore_phrases)
        assert result == text

    def test_none_ignore_phrases_returns_original(self):
        text = "This is my original text."
        result = remove_ignore_phrases(text, "")
        assert result == text

    def test_cleans_extra_whitespace(self):
        text = "Q1: Question text.\n\n\nMy answer here."
        ignore_phrases = "Q1: Question text."
        result = remove_ignore_phrases(text, ignore_phrases)
        assert "Q1: Question text" not in result
        assert "\n\n\n" not in result
        assert "My answer here" in result

    def test_handles_empty_lines_in_ignore_phrases(self):
        text = "Q1: First question. Answer. Q2: Second question. Answer."
        ignore_phrases = "Q1: First question\n\n\nQ2: Second question"
        result = remove_ignore_phrases(text, ignore_phrases)
        assert "Q1: First question" not in result
        assert "Q2: Second question" not in result
        assert "Answer" in result

    def test_case_sensitive_removal(self):
        text = "Q1: Explain the theory. q1: explain the theory."
        ignore_phrases = "Q1: Explain the theory"
        result = remove_ignore_phrases(text, ignore_phrases)
        assert "Q1: Explain the theory" not in result
        assert "q1: explain the theory" in result

    def test_multiple_occurrences_removed(self):
        text = "Instructions: Write in your own words. Paragraph 1. Instructions: Write in your own words. Paragraph 2."
        ignore_phrases = "Instructions: Write in your own words"
        result = remove_ignore_phrases(text, ignore_phrases)
        assert "Instructions: Write in your own words" not in result
        assert "Paragraph 1" in result
        assert "Paragraph 2" in result

    def test_no_match_returns_original(self):
        text = "This is my original text with no matching phrases."
        ignore_phrases = "Q1: Some question\nQ2: Another question"
        result = remove_ignore_phrases(text, ignore_phrases)
        assert result == text


# ---------------------------------------------------------------------------
# clean_text tests
# ---------------------------------------------------------------------------


class TestCleanText:

    def test_collapses_multiple_blank_lines(self):
        text = "Line 1\n\n\n\nLine 2"
        result = clean_text(text)
        assert result == "Line 1\n\nLine 2"

    def test_collapses_multiple_spaces_and_tabs(self):
        text = "Hello     world\t\tPython"
        result = clean_text(text)
        assert result == "Hello world Python"

    def test_replaces_unicode_spaces(self):
        text = "Hello\u00a0World\u200b!"
        result = clean_text(text)
        assert result == "Hello World !"

    def test_removes_spaces_before_newline(self):
        text = "Hello    \nWorld"
        result = clean_text(text)
        assert result == "Hello\nWorld"

    def test_removes_spaces_after_newline(self):
        text = "Hello\n    World"
        result = clean_text(text)
        assert result == "Hello\nWorld"

    def test_strips_leading_and_trailing_whitespace(self):
        text = "   Hello World   \n"
        result = clean_text(text)
        assert result == "Hello World"

    def test_handles_empty_string(self):
        text = ""
        result = clean_text(text)
        assert result == ""

    def test_preserves_normal_text(self):
        text = "This is a normal sentence."
        result = clean_text(text)
        assert result == text

    def test_combines_all_cleaning_steps(self):
        text = "  Hello\t\t\n\n\n  World\u00a0 "
        result = clean_text(text)
        assert result == "Hello\n\nWorld"

    def test_only_whitespace_returns_empty(self):
        text = "   \n\t\n  "
        result = clean_text(text)
        assert result == ""


def test_extract_empty_pdf_gracefully(caplog):
    """Assert that passing an empty/blank PDF returns an empty string gracefully without crashing."""
    import logging

    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.showPage()
    c.save()
    empty_pdf_bytes = buf.getvalue()

    with patch("src.core.document_parser._ocr_pdf_page", return_value=""):
        with caplog.at_level(logging.WARNING):
            result = extract_text_from_pdf(empty_pdf_bytes)

    assert isinstance(result, str)
    assert result.strip() == ""


def test_extract_text_from_doc_success():
    """Test that extract_text_from_doc runs antiword successfully when present."""

    from src.core.document_parser import extract_text_from_doc

    mock_result = MagicMock()
    mock_result.stdout = (
        "This is a test legacy Word Document content extracted by antiword."
    )
    mock_result.returncode = 0

    with patch("shutil.which", return_value="/usr/bin/antiword"):
        with patch("subprocess.run", return_value=mock_result) as mock_run:
            result = extract_text_from_doc(b"fake doc bytes")
            assert (
                result
                == "This is a test legacy Word Document content extracted by antiword."
            )
            mock_run.assert_called_once()
            args, kwargs = mock_run.call_args
            assert args[0][0] == "antiword"
            assert args[0][1].endswith(".doc")


def test_extract_text_from_doc_missing_antiword():
    """Test that extract_text_from_doc raises RuntimeError if antiword is not installed."""
    import pytest

    from src.core.document_parser import extract_text_from_doc

    with patch("shutil.which", return_value=None):
        with pytest.raises(RuntimeError, match="antiword binary is not installed"):
            extract_text_from_doc(b"fake doc bytes")


def test_extract_text_routing_doc():
    """Test that extract_text routes .doc files to extract_text_from_doc."""
    from src.core.document_parser import extract_text

    mock_result = MagicMock()
    mock_result.stdout = "Legacy Word Doc Content"
    mock_result.returncode = 0

    with patch("shutil.which", return_value="/usr/bin/antiword"):
        with patch("subprocess.run", return_value=mock_result):
            result = extract_text(b"\xd0\xcf\x11\xe0fake bytes", "test_file.doc")
            assert result == "Legacy Word Doc Content"


def test_large_pdf_parsing_performance_benchmark():
    """Benchmark test asserting parsing of a 200-page text PDF completes under 3 seconds."""
    import time
    from reportlab.pdfgen import canvas
    
    # 1. Create a 200-page synthetic PDF in-memory using reportlab
    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    for i in range(200):
        # Add enough words per page to bypass OCR (min 8 words)
        c.drawString(100, 750, f"Page {i}: This is a synthetic page of text to parse quickly.")
        c.showPage()
    c.save()
    pdf_bytes = buf.getvalue()
    
    # 2. Time the parsing of the 200-page PDF
    start_time = time.perf_counter()
    parsed_text = extract_text_from_pdf(pdf_bytes)
    duration = time.perf_counter() - start_time
    
    # 3. Assert duration and basic content checks
    assert len(parsed_text) > 0
    assert "Page 199" in parsed_text
    assert duration < 3.0, f"Parsing 200-page PDF took too long: {duration:.2f} seconds (limit: 3.0s)"


def test_extract_text_from_txt_utf16_fallback():
    """Test that extract_text_from_txt successfully decodes a UTF-16 encoded buffer."""
    original_text = "Hello in UTF-16 coding fallback test!"
    utf16_bytes = original_text.encode("utf-16")
    result = extract_text_from_txt(utf16_bytes)
    assert result == original_text


def test_extract_text_from_txt_latin1_fallback():
    """Test that extract_text_from_txt successfully decodes a Latin-1 (ISO-8859-1) encoded buffer."""
    original_text = "Café, naïve, and résumé contents in Latin-1!"
    latin1_bytes = original_text.encode("latin-1")
    result = extract_text_from_txt(latin1_bytes)
    assert result == original_text


def test_extract_text_routing_txt_latin1(tmp_path):
    """Test that extract_text successfully routes and decodes a Latin-1 file."""
    original_text = "Café and naïve text."
    latin1_bytes = original_text.encode("latin-1")
    
    # Write the bytes to a temp file
    file_path = tmp_path / "latin1_test.txt"
    file_path.write_bytes(latin1_bytes)
    
    # Verify routing and decoding
    result = extract_text(str(file_path), "latin1_test.txt")
    assert result == original_text


def test_get_supported_file_extensions():
    """get_supported_file_extensions should return the expected sorted list."""
    from src.core.document_parser import get_supported_file_extensions

    extensions = get_supported_file_extensions()
    assert extensions == [
        ".csv",
        ".docx",
        ".epub",
        ".html",
        ".markdown",
        ".md",
        ".mdown",
        ".pdf",
        ".rtf",
        ".txt",
    ]


class TestCleanWhitespaceOption:
    """Test suite verifying clean_whitespace: bool = True option in extract_text() (#1039)."""

    def test_clean_whitespace_enabled_by_default(self):
        text_with_trailing_spaces = b"Header line   \n\n\n\n\nBody line with spaces   \nFooter line  "
        result = extract_text(text_with_trailing_spaces, "sample.txt")
        assert "Header line\n\nBody line with spaces\nFooter line" in result
        assert "   \n" not in result
        assert "\n\n\n" not in result

    def test_clean_whitespace_explicitly_true(self):
        text_data = b"Line 1    \n\n\n\nLine 2   "
        result = extract_text(text_data, "sample.txt", clean_whitespace=True)
        assert result == "Line 1\n\nLine 2"

    def test_clean_whitespace_disabled(self):
        text_data = b"Line 1    \n\n\n\nLine 2"
        result = extract_text(text_data, "sample.txt", clean_whitespace=False)
        assert "Line 1    " in result
        assert "\n\n\n\nLine 2" in result